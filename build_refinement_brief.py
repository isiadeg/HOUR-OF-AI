from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT = r"C:\Users\assid\HOUR-OF-AI\Hour_of_AI_UI_Refinement_Presentation_Brief.docx"

GREEN = "0B6B3A"
DARK = "163229"
LIGHT_GREEN = "E6F4ED"
PALE = "F4FAF7"
GRAY = "5F6B66"
LIGHT_GRAY = "EEF1EF"
AMBER = "B7791F"
WHITE = "FFFFFF"

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for style_name, size, color, before, after in [
    ("Title", 28, GREEN, 0, 8),
    ("Subtitle", 13, GRAY, 0, 14),
    ("Heading 1", 18, GREEN, 16, 7),
    ("Heading 2", 13.5, DARK, 11, 5),
    ("Heading 3", 11.5, GREEN, 8, 4),
]:
    style = styles[style_name]
    style.font.name = "Aptos Display" if style_name in ("Title", "Heading 1") else "Aptos"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = style_name != "Subtitle"
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_style in ("List Bullet", "List Number"):
    style = styles[list_style]
    style.font.name = "Aptos"
    style.font.size = Pt(10.5)
    style.paragraph_format.left_indent = Inches(0.28)
    style.paragraph_format.first_line_indent = Inches(-0.16)
    style.paragraph_format.space_after = Pt(4)

if "Callout" not in styles:
    callout_style = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
else:
    callout_style = styles["Callout"]
callout_style.font.name = "Aptos"
callout_style.font.size = Pt(11)
callout_style.font.bold = True
callout_style.font.color.rgb = RGBColor.from_string(GREEN)
callout_style.paragraph_format.space_before = Pt(5)
callout_style.paragraph_format.space_after = Pt(5)

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)

def add_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        shade(cell, GREEN)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.bold = True
        r.font.name = "Aptos"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string(WHITE)
    for row_data in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row_data):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)
            if len(table.rows) % 2 == 1:
                shade(cell, PALE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.name = "Aptos"
            r.font.size = Pt(9.2)
            r.font.color.rgb = RGBColor.from_string(DARK)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table

def add_callout(label, text, fill=LIGHT_GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.65)
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    p.style = styles["Callout"]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r2 = p.add_run(text)
    r2.bold = False
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

def add_numbered(items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)

def page_break():
    doc.add_page_break()

# Header/footer
header = section.header
p = header.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("HOUR OF AI  |  PRODUCT UI REFINEMENT")
r.font.name = "Aptos"
r.font.size = Pt(8)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(GRAY)

footer = section.footer
p = footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Nurture Roots - Presentation Brief")
r.font.name = "Aptos"
r.font.size = Pt(8)
r.font.color.rgb = RGBColor.from_string(GRAY)

# Cover
doc.add_paragraph("PRODUCT REFINEMENT BRIEF", style="Subtitle")
title = doc.add_paragraph("Hour of AI Platform", style="Title")
title.paragraph_format.space_before = Pt(70)
sub = doc.add_paragraph("Clarifying Sessions, Classes, Programmes and Role-Based Actions", style="Subtitle")
sub.paragraph_format.space_after = Pt(24)
add_callout(
    "Core outcome",
    "The interface now distinguishes what users are viewing from what they are allowed to manage. Sessions represent dated events, classes represent recurring learning groups, and programmes represent a different relationship for parents, volunteers and hubs."
)
doc.add_paragraph("Prepared for product presentation and stakeholder review", style="Subtitle")
doc.add_paragraph("June 2026", style="Subtitle")

page_break()

doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "The original parent and hub dashboards reused labels, routes and actions across concepts that were operationally different. "
    "This made the interface look complete while leaving key questions unanswered: Was a user viewing a class or a session? "
    "What did enrolment mean for a hub? Could a volunteer approve themselves for a programme? Which child was actually enrolled?"
)
doc.add_paragraph(
    "The refinement work reorganised the experience around three product principles:"
)
add_numbered([
    "Use routes and labels that match the object being viewed.",
    "Separate recurring structures (classes) from dated events (sessions).",
    "Make programme actions role-specific instead of treating every account as generically enrolled."
])
add_callout(
    "Presentation thesis",
    "The work is not simply a visual redesign. It is a correction of the platform's information architecture and role permissions."
)

doc.add_heading("What Was Implemented", level=2)
add_bullets([
    "Session-first routes and labels for parent and hub schedules.",
    "Backward-compatible redirects from older /classes schedule URLs.",
    "Class-level cards and details for hub administrators, aligned with volunteer class information.",
    "Overview links that open a filtered and highlighted session instead of repeating a thin details page.",
    "Distinct Programme dashboards for parents, volunteers and hub administrators.",
    "An editable hub operational profile used as readiness context for programme requests."
])

doc.add_heading("2. The Structural Problem", level=1)
add_table(
    ["Original issue", "Why it was confusing", "Refined model"],
    [
        ("Schedule pages used /classes", "The page primarily displayed dated sessions rather than reusable class records.", "Use /sessions for dated schedules and preserve /classes for class records."),
        ("Hub class cards opened session details", "A recurring class card led to one occurrence, mixing two levels of information.", "Class cards open class-level details; session lists handle dated occurrences."),
        ("One generic Programmes page", "Parents, volunteers and hubs appeared to perform the same enrolment action.", "Each role manages a different relationship with a programme."),
        ("Hub registration data was static later", "Infrastructure and schedule could change after onboarding, but the dashboard could not update them.", "Hub Profile owns editable operational information used by readiness checks.")
    ],
    widths=[1.55, 2.5, 2.6]
)

page_break()

doc.add_heading("3. Sessions and Classes", level=1)
doc.add_heading("Decision", level=2)
doc.add_paragraph(
    "A class is a recurring learning group. A session is a dated occurrence of that class. The routes, cards and navigation should preserve that distinction."
)
add_table(
    ["Concept", "Meaning", "Example route", "Typical information"],
    [
        ("Class", "Recurring group or assignment", "/hub/bright-stars/classes/hc1", "Programme, age group, volunteers, learners, recurring schedule, history"),
        ("Session", "Specific dated event", "/hub/bright-stars/sessions", "Date, time, status, join link, session actions"),
        ("Focused session", "A session selected from Overview", "/hub/bright-stars/sessions?session=hcs1", "Filtered list with the matching session highlighted")
    ],
    widths=[1.0, 1.65, 2.15, 1.85]
)

doc.add_heading("Implemented Routing and Navigation", level=2)
add_bullets([
    "Parent schedule changed from /parent/classes to /parent/sessions.",
    "Hub schedule changed from /hub/{hub}/classes to /hub/{hub}/sessions.",
    "Older schedule URLs redirect to the new canonical routes.",
    "Session-detail links remain available for richer future operational records.",
    "Back links now explicitly return to Sessions, category Classes, Hub Learners or My Learners as appropriate."
])

doc.add_heading("Overview Card Behaviour", level=2)
doc.add_paragraph(
    "The Next Session card previously used a generic View Details action. Because the individual session page repeated information already visible on the list card, the action now opens the Sessions page with the relevant session filtered and highlighted."
)
add_callout(
    "Why this is better",
    "It preserves context, avoids an unnecessary page transition, creates a shareable URL, and still gives users a clear View all sessions reset."
)

doc.add_heading("4. Hub Class Experience", level=1)
doc.add_paragraph(
    "Hub administrators need class-level visibility comparable to volunteers because they coordinate learners, facilities and delivery. "
    "The hub class cards were therefore changed from dated-session cards into recurring class cards."
)
add_table(
    ["Before", "After"],
    [
        ("Card showed 'Sat, 12 Jul 2025 - 10:00 AM'.", "Card shows 'Every Saturday, 10:00 AM'."),
        ("Details opened a session occurrence.", "Details opens /hub/{hub}/classes/{classId}."),
        ("Limited operational context.", "Class detail includes age group, delivery, lead and backup volunteers, learners, recurring schedule and session history."),
        ("Class and session labels were mixed.", "Class pages retain class language; schedule pages use session language.")
    ],
    widths=[3.25, 3.4]
)

page_break()

doc.add_heading("5. Programme Model by Role", level=1)
add_callout(
    "Guiding principle",
    "Registration captures initial interest and baseline information. The Programme dashboard manages the user's ongoing relationship with each programme."
)
add_table(
    ["Role", "What registration means", "What the dashboard manages", "Operational result"],
    [
        ("Parent", "Initial programme interest for the family", "Learner-by-learner enrolment eligibility and requests", "A child may become enrolled and later assigned to a class"),
        ("Volunteer", "Programmes they would like to support", "Interest, training, approval and matching eligibility", "An approved volunteer becomes eligible for class assignment"),
        ("Hub admin", "Programmes the hub intends to run", "Approval, readiness, activation and capacity", "An approved and ready programme becomes active at the hub"),
        ("Administrator", "Not an enrolment role", "Review, approval, capacity and matching", "Relationships become operational after checks")
    ],
    widths=[0.85, 1.65, 2.25, 1.9]
)

doc.add_heading("Parent Programme Dashboard", level=2)
add_bullets([
    "A programme expands to show every learner attached to the parent account.",
    "Each learner has a status: Enrolled, Eligible, Pending, Waitlisted, Not eligible or Removal requested.",
    "Parents select one or more eligible learners before requesting enrolment.",
    "Enrolled learners can have a removal request submitted.",
    "This replaces the inaccurate idea that the parent account itself is enrolled."
])

doc.add_heading("Volunteer Programme Dashboard", level=2)
add_bullets([
    "Eligible: training and safeguarding requirements are complete.",
    "Training required: interest exists, but programme requirements are incomplete.",
    "Under review: interest has been submitted and approval is progressing.",
    "Available: the volunteer may register interest.",
    "Approval does not assign a class; confirmed assignments remain under My Classes."
])

doc.add_heading("Hub Programme Dashboard", level=2)
add_bullets([
    "Active: currently running at the hub.",
    "Approved: permitted but not yet launched.",
    "Pending review: a programme request is being assessed.",
    "Setup required: operational requirements remain incomplete.",
    "Available: the hub can submit a programme request."
])

page_break()

doc.add_heading("6. Hub Profile and Programme Requests", level=1)
doc.add_heading("The Duplication Problem", level=2)
doc.add_paragraph(
    "Hub onboarding already collects programme choices, infrastructure, session preferences and support needs. "
    "Repeating all of these fields during every programme request would create friction and inconsistent data."
)

doc.add_heading("Refined Ownership", level=2)
add_table(
    ["Hub Profile owns", "Programme request owns"],
    [
        ("Internet reliability and power supply", "Expected learners for this programme"),
        ("Number of available computers/devices", "Target age group"),
        ("Overall learner capacity", "Preferred programme start month"),
        ("Supported age groups", "Programme-specific schedule"),
        ("General operating days and time slots", "Delivery preference"),
        ("Timezone and ongoing support needs", "Programme-specific support or accessibility needs")
    ],
    widths=[3.25, 3.4]
)

doc.add_heading("Implemented Hub Operations Editor", level=2)
add_bullets([
    "Hub administrators can update internet, computers, power, learner capacity and timezone.",
    "They can add or remove age groups, operating days, time slots and support needs.",
    "Changes persist locally in the current prototype.",
    "Programme requests display a read-only readiness summary sourced from Hub Profile."
])

add_callout(
    "Practical example",
    "If a hub adds solar power, acquires more computers or starts opening on Sundays, the administrator updates Hub Profile once. New programme requests immediately use the revised readiness context."
)

doc.add_heading("7. User Journeys After Refinement", level=1)
doc.add_heading("Parent requests a programme", level=2)
add_numbered([
    "Open Programmes and select a programme.",
    "Review each learner's current eligibility and status.",
    "Select one or more eligible learners.",
    "Submit an enrolment request.",
    "Track pending, waitlisted or enrolled status."
])

doc.add_heading("Volunteer becomes eligible", level=2)
add_numbered([
    "Open Programmes and register interest.",
    "Complete programme-specific training and safeguarding requirements.",
    "Track the approval review.",
    "Become eligible for matching.",
    "Receive a confirmed class assignment under My Classes."
])

doc.add_heading("Hub requests a programme", level=2)
add_numbered([
    "Keep Hub Profile operational information current.",
    "Open Programmes and select an available programme.",
    "Review the readiness summary sourced from Hub Profile.",
    "Provide only programme-specific learner, age, schedule and support information.",
    "Track review, setup, approval and activation."
])

page_break()

doc.add_heading("8. Prototype Status and Backend Requirements", level=1)
doc.add_paragraph(
    "The current implementation establishes the front-end workflows and expected data relationships. Mock state is intentionally used so the product logic can be reviewed before backend contracts are finalised."
)
add_table(
    ["Implemented in UI", "Required from backend"],
    [
        ("Role-specific Programme pages and actions", "Persistent role-programme relationships"),
        ("Learner eligibility and request states", "Real age, prerequisite, geography and capacity checks"),
        ("Volunteer training and approval states", "Training completion, safeguarding verification and approval workflow"),
        ("Hub readiness summary and editable operations", "Persisted hub profile, audit history and administrator review"),
        ("Focused session query links", "Session API filters and stable session identifiers"),
        ("Mock request and withdrawal actions", "Notifications, approvals, waitlists and status transitions")
    ],
    widths=[3.15, 3.5]
)

doc.add_heading("9. Key Presentation Messages", level=1)
add_bullets([
    "The redesign corrects the product model before backend complexity is added.",
    "Classes and sessions are no longer treated as interchangeable.",
    "Programme participation is now expressed at the correct level: child, volunteer eligibility or hub activation.",
    "Registration remains the starting point; dashboards support change over time.",
    "Hub operational data has one editable source of truth instead of being repeatedly collected.",
    "The prototype now exposes the backend contracts the product will need."
])

doc.add_heading("10. Recommended Next Steps", level=1)
add_numbered([
    "Define API contracts for learner-programme, volunteer-programme and hub-programme statuses.",
    "Connect Hub Profile operations to persistent backend storage and administrator review.",
    "Add administrator queues for enrolment, volunteer approval and hub activation requests.",
    "Replace mock eligibility with real capacity, age and prerequisite checks.",
    "Add notifications and audit history for request status changes.",
    "Run usability testing with one parent, one volunteer and one hub administrator journey."
])

add_callout(
    "Final position",
    "The interface is now structured around real operational relationships. The next phase is to make those relationships persistent, reviewable and automated."
)

# Keep tables and rows readable
for table in doc.tables:
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.keep_together = True

doc.core_properties.title = "Hour of AI UI Refinement Presentation Brief"
doc.core_properties.subject = "Product structure, routing, role-based programmes and hub operations"
doc.core_properties.author = "Nurture Roots Product Team"
doc.save(OUT)
print(OUT)
